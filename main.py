import os
import streamlit as st
import subprocess
import zipfile
from docx import Document
import io
from langdetect import detect
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain

# Configuración de Streamlit
st.set_page_config('TranscriptorAV')
st.title("TranscriptorAV: Tu asistente para transcripciones de audio y video")

# Idioma por defecto
lang = 'es'

# Obtener y manejar la clave API de OpenAI
API_KEY = os.environ.get('API_KEY')
if not API_KEY:
    try:
        with open('api_key.txt', 'r') as f:
            API_KEY = f.read().strip()
    except FileNotFoundError:
        with open('api_key.txt', 'w') as f:
            f.write('tu_clave_de_api')
        API_KEY = st.text_input('OpenAI API Key', type='password')

language_choice = st.sidebar.selectbox(
    "Selecciona el idioma",
    ("Spanish", "English")
)

# Set the language based on the user's selection
lang = 'es' if language_choice == "Spanish" else 'en'

# Sidebar configurations
model_choice = st.sidebar.selectbox(
    "Selecciona el modelo de Whisper",
    ("tiny", "base", "small", "medium", "large")
)

# Add option to ask questions to the transcribed text
operation_choice = st.sidebar.selectbox(
    "Qué deseas hacer?",
    ("Transcribir y Descargar", "Preguntar al Audio/Video")
)

# Footer
st.sidebar.markdown('---')
st.sidebar.subheader('Creado por:')
st.sidebar.markdown('Alexander Oviedo Fadul')
st.sidebar.markdown(
    "[GitHub](https://github.com/bladealex9848) | [Website](https://alexander.oviedo.isabellaea.com/) | [Instagram](https://www.instagram.com/alexander.oviedo.fadul) | [Twitter](https://twitter.com/alexanderofadul) | [Facebook](https://www.facebook.com/alexanderof/) | [WhatsApp](https://api.whatsapp.com/send?phone=573015930519&text=Hola%20!Quiero%20conversar%20contigo!%20)"
)

# Función para crear embeddings a partir de texto transcribido


@st.cache_resource
def create_embeddings_from_text(text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        length_function=len
    )
    chunks = text_splitter.split_text(text)

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
    knowledge_base = FAISS.from_texts(chunks, embeddings)

    return knowledge_base, text


@st.cache_resource
def process_file(uploaded_file):
    file_name_without_extension = os.path.splitext(uploaded_file.name)[0]
    temp_file_path = f"{file_name_without_extension}"
    with open(temp_file_path, "wb") as f:
        f.write(uploaded_file.getvalue())

    cmd = f'whisper "{temp_file_path}" --model {model_choice} --language {language_choice}'
    subprocess.run(cmd, shell=True)

    doc_path = f"{file_name_without_extension}.docx"
    doc = Document()
    doc.add_heading('Transcripción', 0)

    file_path = f"{file_name_without_extension}.txt"
    with open(file_path, "r") as txt_file:
        transcription = txt_file.read()

    doc.add_paragraph(transcription)
    doc.save(doc_path)

    zip_file_name = f"{file_name_without_extension}.zip"
    with zipfile.ZipFile(zip_file_name, 'w') as zipf:
        for ext in ["json", "srt", "tsv", "txt", "vtt", "docx"]:
            zipf.write(f"{file_name_without_extension}.{ext}")

    with open(zip_file_name, 'rb') as f:
        zip_bytes = io.BytesIO(f.read())
        st.download_button('Descargar archivo comprimido',
                           zip_bytes, zip_file_name)

    # Cleanup
    os.remove(temp_file_path)
    for ext in ["json", "srt", "tsv", "txt", "vtt", "docx"]:
        os.remove(f"{file_name_without_extension}.{ext}")

    knowledge_base, _ = create_embeddings_from_text(transcription)

    return transcription, knowledge_base


def ask_transcription(transcription, knowledge_base):
    st.header("Preguntar al Audio/Video")
    user_question = st.text_input("Haz una pregunta sobre la transcripción:")
    if user_question:
        os.environ["OPENAI_API_KEY"] = API_KEY
        docs = knowledge_base.similarity_search(user_question, 10)
        llm = ChatOpenAI(model_name='gpt-3.5-turbo')
        chain = load_qa_chain(llm, chain_type="stuff")
        respuesta = chain.run(input_documents=docs, question=user_question)
        st.write(respuesta)


def main():
    st.write("Sube tu archivo de audio o video.")
    uploaded_file = st.file_uploader("Elige un archivo")
    if uploaded_file:
        st.write("Archivo cargado:", uploaded_file.name)
        if operation_choice == "Transcribir y Descargar":
            if st.button("Procesar"):
                process_file(uploaded_file)
                st.success("Archivo procesado exitosamente.")
        elif operation_choice == "Preguntar al Audio/Video":
            transcription, knowledge_base = process_file(uploaded_file)
            ask_transcription(transcription, knowledge_base)


if __name__ == "__main__":
    main()
